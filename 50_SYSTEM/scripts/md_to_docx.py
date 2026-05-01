import sys
import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def md_to_docx(md_path, docx_path):
    print(f"🚀 Leyendo {md_path}...")
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    
    # Estilo base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    for line in lines:
        line = line.strip()
        if not line or line.startswith("---"):
            continue
        
        # Headings
        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        
        # Tables (Very basic markdown table parser)
        elif "|" in line and "---" not in line:
            cells = [c.strip() for c in line.split("|") if c.strip()]
            if not any(isinstance(child, str) for child in cells): continue # skip empty
            # If it's a new table block... 
            # (Simplification: we'll just treat lines with | as text for now to avoid complexity 
            # unless a table helper is needed. Let's try to add proper tables.)
            table = doc.add_table(rows=1, cols=len(cells))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, text in enumerate(cells):
                hdr_cells[i].text = text
        
        # Lists
        elif line.startswith("- [ ]") or line.startswith("- [x]") or line.startswith("- "):
            clean_line = re.sub(r'^- \[[ x]\]\s*', '', line).lstrip("- ")
            doc.add_paragraph(clean_line, style='List Bullet')
        
        # Standard paragraph
        else:
            # Simple bold/italic removal for clean Docx
            clean_text = line.replace("**", "").replace("__", "").replace("*", "")
            doc.add_paragraph(clean_text)

    print(f"✅ Guardando en {docx_path}...")
    doc.save(docx_path)

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Uso: python md_to_docx.py input.md output.docx")
    else:
        md_to_docx(sys.argv[1], sys.argv[2])
